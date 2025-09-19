from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
import uvicorn
import os
import json
import tempfile
from datetime import datetime
from typing import Optional, List
import PyPDF2
import docx
from PIL import Image
import io
from sqlalchemy import text
import time
import hashlib

# Import database components
from database import get_db, init_db, User, Document, Conversation, Analytics

app = FastAPI(
    title="EduVision AI Backend",
    description="Advanced AI-powered educational platform backend with multimodal capabilities",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize database on startup
@app.on_event("startup")
async def startup_event():
    init_db()

# Helper functions for AI processing (simplified for demo)
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        return f"Error extracting PDF text: {str(e)}"

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        return f"Error extracting DOCX text: {str(e)}"

def analyze_image(image_content: bytes) -> str:
    """Analyze image content (simplified)"""
    try:
        image = Image.open(io.BytesIO(image_content))
        # Simplified analysis - in production, use actual AI models
        return f"Image analysis: This appears to be a {image.format} image with dimensions {image.size[0]}x{image.size[1]} pixels. The image contains visual content that could include diagrams, charts, text, or other educational materials."
    except Exception as e:
        return f"Error analyzing image: {str(e)}"

def generate_summary(text: str, max_length: int = 150) -> str:
    """Generate text summary"""
    if not text or len(text.strip()) == 0:
        return "No text provided for summarization."
    
    # Clean and prepare text
    text = text.strip()
    
    # If text is already short, return as is
    if len(text) <= max_length:
        return text
    
    # Split into sentences
    import re
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
    
    if len(sentences) <= 1:
        # If only one sentence, truncate it
        return text[:max_length-3] + "..." if len(text) > max_length else text
    
    # Extractive summarization: select key sentences
    summary_sentences = []
    current_length = 0
    
    # Always include first sentence
    if sentences:
        summary_sentences.append(sentences[0])
        current_length = len(sentences[0])
    
    # Add middle and last sentences if they fit
    if len(sentences) > 2:
        middle_idx = len(sentences) // 2
        last_idx = len(sentences) - 1
        
        for idx in [middle_idx, last_idx]:
            sentence = sentences[idx]
            if current_length + len(sentence) + 2 <= max_length:  # +2 for ". "
                summary_sentences.append(sentence)
                current_length += len(sentence) + 2
    
    summary = '. '.join(summary_sentences)
    if not summary.endswith('.'):
        summary += '.'
    
    # Final length check
    if len(summary) > max_length:
        summary = summary[:max_length-3] + "..."
    
    return summary

def answer_question(context: str, question: str) -> str:
    """Answer question based on context"""
    if not context or not question:
        return "Please provide both context and a question."
    
    context_lower = context.lower()
    question_lower = question.lower()
    
    # Simple keyword-based answering (in production, use AI models)
    if any(word in question_lower for word in ['what', 'define', 'definition']):
        # Look for definitions or explanations
        sentences = context.split('.')
        for sentence in sentences:
            if any(word in sentence.lower() for word in question_lower.split() if len(word) > 3):
                return sentence.strip() + "."
    
    if any(word in question_lower for word in ['how', 'process', 'steps']):
        # Look for process descriptions
        sentences = context.split('.')
        process_sentences = []
        for sentence in sentences:
            if any(word in sentence.lower() for word in ['first', 'then', 'next', 'step', 'process']):
                process_sentences.append(sentence.strip())
        if process_sentences:
            return '. '.join(process_sentences[:3]) + "."
    
    # Default: return relevant sentences
    sentences = context.split('.')
    relevant_sentences = []
    question_words = [word for word in question_lower.split() if len(word) > 3]
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        if any(word in sentence_lower for word in question_words):
            relevant_sentences.append(sentence.strip())
    
    if relevant_sentences:
        return '. '.join(relevant_sentences[:2]) + "."
    else:
        return "I couldn't find a specific answer to your question in the provided context. Please try rephrasing your question or provide more relevant content."

# Pydantic models
class QuestionRequest(BaseModel):
    question: str
    document_id: Optional[str] = None

class SummarizationRequest(BaseModel):
    text: str
    max_length: Optional[int] = 150

class TextToSpeechRequest(BaseModel):
    text: str
    voice: Optional[str] = "default"

class SimilarityRequest(BaseModel):
    text1: str
    text2: str

# Health check endpoint
@app.get("/")
async def root():
    return {"message": "EduVision AI Backend is running!", "status": "healthy"}

@app.get("/health")
async def health_check(db: Session = Depends(get_db)):
    try:
        # Test database connection
        db.execute(text("SELECT 1"))
        return {
            "status": "healthy", 
            "service": "EduVision AI Backend",
            "database": "connected",
            "timestamp": datetime.utcnow().isoformat()
        }
    except Exception as e:
        return {
            "status": "unhealthy", 
            "service": "EduVision AI Backend",
            "database": "disconnected",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        }

# Document Question Answering
@app.post("/document-qa")
async def document_qa(
    document_id: int = Form(...),
    question: str = Form(...),
    db: Session = Depends(get_db)
):
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Generate answer
        answer = answer_question(document.content, question)
        
        # Save conversation
        conversation = Conversation(
            user_id=1,  # Demo user
            document_id=document_id,
            question=question,
            answer=answer,
            service_type="document_qa"
        )
        db.add(conversation)
        db.commit()
        
        return {
            "answer": answer,
            "document_title": document.filename,
            "confidence": "High" if len(answer) > 50 else "Medium"
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Visual Question Answering
@app.post("/visual-qa")
async def visual_qa(
    image: UploadFile = File(...),
    question: str = Form(...),
    db: Session = Depends(get_db)
):
    try:
        # Read and analyze image
        image_content = await image.read()
        image_analysis = analyze_image(image_content)
        
        # Generate contextual answer based on image analysis and question
        question_lower = question.lower()
        
        # Analyze question intent
        if any(word in question_lower for word in ['what', 'describe', 'see', 'show', 'content']):
            # Descriptive question - focus on image content
            answer = f"Based on my analysis of this image:\n\n{image_analysis}\n\nTo answer your question '{question}': The image contains visual elements that I've analyzed above. The technical and visual characteristics suggest this could be useful for educational purposes."
            
        elif any(word in question_lower for word in ['how', 'why', 'explain', 'meaning']):
            # Explanatory question - provide educational context
            answer = f"Here's my analysis to help explain what you're seeing:\n\n{image_analysis}\n\nRegarding '{question}': Based on the visual characteristics, this image appears to be educational content. The analysis above provides insights into its structure and potential use cases."
            
        elif any(word in question_lower for word in ['color', 'colours', 'bright', 'dark']):
            # Color-related question - focus on visual characteristics
            answer = f"Focusing on the visual aspects you asked about:\n\n{image_analysis}\n\nFor your question '{question}': The color and brightness analysis above should help answer your specific inquiry about the visual properties of this image."
            
        elif any(word in question_lower for word in ['size', 'dimension', 'resolution', 'quality']):
            # Technical question - focus on technical details
            answer = f"Here are the technical specifications:\n\n{image_analysis}\n\nRegarding '{question}': The technical details section above provides the specific information you're looking for about the image properties."
            
        else:
            # General question - provide comprehensive analysis
            answer = f"Here's a comprehensive analysis of your image:\n\n{image_analysis}\n\nFor your question '{question}': This analysis should provide the context needed to understand the image content and its potential educational applications."
        
        # Save conversation
        conversation = Conversation(
            user_id=1,  # Demo user
            question=question,
            answer=answer,
            service_type="visual_qa"
        )
        db.add(conversation)
        db.commit()
        
        return {
            "answer": answer,
            "image_analysis": image_analysis,
            "confidence": "Medium"
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Text Summarization
@app.post("/summarize")
async def summarize_text(
    text: str = Form(...),
    max_length: int = Form(150),
    db: Session = Depends(get_db)
):
    try:
        if not text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        # Generate summary
        summary = generate_summary(text, max_length)
        
        # Save conversation
        conversation = Conversation(
            user_id=1,  # Demo user
            question=f"Summarize text (max {max_length} chars)",
            answer=summary,
            service_type="summarization"
        )
        db.add(conversation)
        db.commit()
        
        return {
            "summary": summary,
            "original_length": len(text),
            "summary_length": len(summary),
            "compression_ratio": f"{(1 - len(summary)/len(text))*100:.1f}%"
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Text-to-Speech
@app.post("/text-to-speech")
async def text_to_speech(
    text: str = Form(...),
    voice: str = Form("default"),
    db: Session = Depends(get_db)
):
    try:
        if not text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        # Simulate TTS processing
        word_count = len(text.split())
        estimated_duration = word_count * 0.5  # Rough estimate: 2 words per second
        
        # Save conversation
        conversation = Conversation(
            user_id=1,  # Demo user
            question=f"Convert to speech using {voice} voice",
            answer=f"Text-to-speech conversion completed for {word_count} words",
            service_type="text_to_speech"
        )
        db.add(conversation)
        db.commit()
        
        return {
            "message": f"Text successfully converted to speech using {voice} voice!",
            "word_count": word_count,
            "estimated_duration": f"{estimated_duration:.1f} seconds",
            "voice_used": voice,
            "audio_url": "/audio/placeholder.mp3"  # Placeholder
        }
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# File Upload
@app.post("/upload")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """
    Upload and process educational documents
    """
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        extracted_text = ""
        if file.filename.lower().endswith('.pdf'):
            extracted_text = extract_text_from_pdf(content)
        elif file.filename.lower().endswith('.docx'):
            extracted_text = extract_text_from_docx(content)
        elif file.filename.lower().endswith(('.txt', '.md')):
            extracted_text = content.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        # Save to database
        document = Document(
            filename=file.filename,
            content=extracted_text,
            file_type=file.content_type,
            user_id=1  # Demo user
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        return {
            "message": f"File '{file.filename}' uploaded and processed successfully!",
            "document_id": document.id,
            "extracted_text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text,
            "total_characters": len(extracted_text)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

# Analytics endpoint
@app.get("/analytics")
async def get_analytics(db: Session = Depends(get_db)):
    """
    Get platform usage analytics
    """
    try:
        # Get conversation counts by feature
        conversations = db.query(Conversation).all()
        
        service_counts = {}
        for conv in conversations:
            service_counts[conv.service_type] = service_counts.get(conv.service_type, 0) + 1
        
        # Get document count
        document_count = db.query(Document).count()
        
        analytics = {
            "total_users": 1,  # Demo user
            "documents_processed": document_count,
            "questions_answered": len(conversations),
            "service_usage": service_counts,
            "database_status": "connected",
            "last_updated": datetime.now().isoformat()
        }
        return analytics
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Get conversation history
@app.get("/conversations")
async def get_conversations(limit: int = 10, db: Session = Depends(get_db)):
    """
    Get recent conversation history
    """
    try:
        conversations = db.query(Conversation).order_by(Conversation.created_at.desc()).limit(limit).all()
        return [
            {
                "id": conv.id,
                "service_type": conv.service_type,
                "question": conv.question,
                "answer": conv.answer,
                "timestamp": conv.created_at.isoformat()
            }
            for conv in conversations
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
